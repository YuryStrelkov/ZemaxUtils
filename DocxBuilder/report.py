import os.path
import numpy as np
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_LINE_SPACING
from docx.shared import Inches, Cm
from typing import Union, Tuple
from docx.shared import Pt
from docx import Document
from matplotlib import pyplot as plt
from docx.enum.section import WD_ORIENTATION
from Geometry import poly_regression, linear_regression, Vector2
from ResultBuilder import ResultFile, draw_spot, ResultVisualSettings, draw_pos_from_angle, draw_psf, \
    draw_psf_cross_section, draw_mtf
from docx.enum.section import WD_SECTION_START, WD_ORIENTATION


class Report:
    @property
    def images_count(self) -> int:
        return self._images_count

    @images_count.setter
    def images_count(self, value: int) -> None:
        assert isinstance(value, int)
        self._images_count = max(0, value)

    @property
    def tables_count(self) -> int:
        return self._images_count

    @tables_count.setter
    def tables_count(self, value: int) -> None:
        assert isinstance(value, int)
        self._tables_count = max(0, value)

    def __init__(self):
        self._images_count: int = 0
        self._tables_count: int = 0
        self._formulae_count: int = 0
        self._docx: Document = Document(r'E:\Github\ZemaxUtils\ZemaxUtils\DocxBuilder\report_template.docx')
        self._report_tmp: str = ''

    def keep_table_on_one_page(self):
        tags = self._docx.element.xpath('//w:tr[position() < last()]/w:tc/w:p')
        for tag in tags:
            ppr = tag.get_or_add_pPr()
            ppr.keepNext_val = True

    def add_paragraph(self, text: str = 'Paragraph', font_name: str = 'Times New Roman', font_size: int = 14,
                      alignment=WD_PARAGRAPH_ALIGNMENT.CENTER, italic: bool = False, bold: bool = False,
                      indent: float = 1.5):
        p = self._docx.add_paragraph()
        p.paragraph_format.line_spacing = indent
        p.paragraph_format.alignment = alignment
        t = p.add_run(text)
        t.font.name = font_name
        t.font.size = Pt(font_size)
        t.font.bold = bold
        t.italic = italic
        return self

    def add_image(self, src: str, desc: str = '',
                  alignment=WD_PARAGRAPH_ALIGNMENT.CENTER,
                  size: Union[Tuple[float, float], float] = 17.0, font_size: int = 14):

        self._images_count += 1
        try:
            if isinstance(size, float) or isinstance(size, int):
                img = self._docx.add_picture(src, width=Cm(size))
            if isinstance(size, Tuple):
                img = self._docx.add_picture(src, width=Cm(size[0]), height=Cm(size[1]))
        except Exception as _:
            img = self._docx.add_picture(src, width=Cm(12.0))

        last_paragraph = self._docx.paragraphs[-1]
        last_paragraph.alignment = alignment
        if desc == "":
            f_name = '.'.join(v for v in src.split("\\")[-1].split(".")[:-1])
            if f_name.startswith("psf"):
                desc = f"Результат моделирования PSF по всем полям для длины волны {f_name.split(' ')[-1]} мкм."
            elif f_name.startswith("psf_cross"):
                desc = f"Результат моделирования сечения PSF плоскостями XoZ и YoZ" \
                              f" по всем длинам волн для поля с номером {f_name.split(' ')[-1]}."
            elif f_name.startswith("mtf"):
                desc = f"Результат моделирования MTF по всем длинам волн для всех полей."
            elif f_name.startswith("spot"):
                desc = f"Результат моделирования SPOT диаграммы по всем длинам волн для всех полей."
            else:
                desc = "Нет описания к рисунку."

        self.add_paragraph(f"Рисунок {self._images_count}. {desc}\n", font_size=font_size, alignment=alignment,
                           italic=False, indent=1.0)

    def add_table(self, headers, data, description: str = "", font_name: str = 'Times New Roman', font_size: int = 14,
                  alignment=WD_PARAGRAPH_ALIGNMENT.CENTER, italic: bool = False, format_provider=lambda s: str(s)):
        self._tables_count += 1

        self.add_paragraph(text=f"\nТаблица. {self._tables_count}. "
                                f"{'Нет описания к таблице' if description == '' else description}\n",
                           italic=italic, font_name=font_name, font_size=14, indent=1.0)

        table = self._docx.add_table(rows=len(data) // len(headers) + 1, cols=len(headers))
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = 'Table Grid'
        header_cells = table.rows[0].cells
        for row in table.rows:
            row.height = Cm(0.8)

        for cell, header in zip(header_cells, headers):
            cell.text = str(header)
            font = cell.paragraphs[0].runs[0].font
            font.bold = True
            font.name = font_name
            font.size = Pt(font_size)
            cell.paragraphs[0].alignment = alignment
        row, col = 0, 0
        try:
            for record_index, record in enumerate(data):
                row, col = divmod(record_index, len(headers))
                # print(f'row: {row:>3} | col: {col:>3}')
                cell = table.rows[row + 1].cells[col]
                cell.text = format_provider(record)
                font = cell.paragraphs[0].runs[0].font
                font.name = font_name
                font.size = Pt(font_size)
        except IndexError as er:
            print(f'IndexError at :: row: {row:>3} | col: {col:>3}\n {er.args}' )
        self._docx.add_page_break()

    def save(self, file_path: str):
        self._docx.save(file_path if file_path.endswith('docx') else file_path + '.docx')

    def _basic_params(self, file: ResultFile):
        self.add_paragraph(f"\tОсновные параметры схемы и окружения представлены в таблице {self._tables_count + 1}.",
                           alignment=WD_PARAGRAPH_ALIGNMENT.JUSTIFY)
        self.add_table(description='Основные параметры исследуемой схемы',
                       headers=('Параметр схемы', 'Значение'),
                       data=("количество элементов",                      str(file.entries_count),
                             "апертура, [мм]",                            str(file.aperture_value),
                             "температура окр. среды (С)",                str(file.temp_c),
                             "давление окр. среды (MPa)",                 str(file.pressure_atm),
                             "эффективное фокусное расстояние, [мм]",     str(file.efl),
                             "входной диаметр зрачка, [мм]",              str(file.entrance_pupil_dia),
                             "входное положение зрачка, [мм]",            str(file.entrance_pupil_pos),
                             "выходной диаметр зрачка, [мм]",             str(file.exit_pupil_dia),
                             "выходное положение зрачка, [мм]",           str(file.exit_pupil_pos),
                             "параксиальная высота изображения, [мм]",    str(file.parax_image_height),
                             "параксиальное увеличение",                  str(file.parax_magnification),
                             "угловое увеличение",                        str(file.angular_magnification),
                             "полный опт. путь, [мм]",                    str(file.total_track),
                             "сдвиг зрачка по x, [мм]",                   str(file.x_pupil_shift),
                             "сдвиг зрачка по y, [мм]",                   str(file.y_pupil_shift),
                             "сдвиг зрачка по z, [мм]",                   str(file.z_pupil_shift),
                             "номер STOP поверхности",                    str(file.stop_surface_number),
                             "Интегральный коэффициент пропускания, [%]", str(70.0),
                             'Коэффициент виньетирования, [%]',           str(30.6),
                             'Коэффициент экранирования, [%]',            str(32.2)))

        self.add_paragraph(f"\n\tСписок параметров полей схемы представлен в таблице {self._tables_count + 1}.",
                           alignment=WD_PARAGRAPH_ALIGNMENT.JUSTIFY)
        data = []
        for v in file.fields.values():
            for vi in v:
                data.append(f'{vi: .4}')
        self.add_table(description='Параметры полей в исследуемой схеме.',
                       headers=("FLDX", "FLDY", "FWGT", "FVDX", "FVDY", "FVCX", "FVCY", "FVAN"), data=data)
        data.clear()
        for wl, ww in zip(file.wavelengths, file.wavelengths_weights):
            data.append(f'{wl: .4}')
            data.append(f'{ww: .4}')
        self.add_table(description='Параметры длин волн в исследуемой схеме.', headers=("WL,[мкм]", "Вес"), data=data)

    def _spot_diagrams(self, file: ResultFile):
        visual_settings = ResultVisualSettings()
        visual_settings.subplots = (2, 2)
        visual_settings.width = 17
        visual_settings.height = 17
        visual_settings.font_size = 12
        visual_settings.bounds = (0.05, 0.05, 0.95, 0.95)
        visual_settings.h_space = 0.125
        visual_settings.w_space = 0.125
        angles = ",".join(f" {{ x = {file.fields[fid+1].FLDX:.3}, y = {file.fields[fid+1].FLDY:.3} }}"
                          for fid in range(visual_settings.subplots_count))
        self.add_paragraph(f"\tSpot - диаграмма. Иллюстрирует распределение точек пересечения лучей с плоскостью "
                           f"изображения. Примеры для полей схемы c углами падения: {angles}, "
                           f"представлены на рисунках {self._images_count + 1}.\n",
                           alignment=WD_PARAGRAPH_ALIGNMENT.JUSTIFY)
        fig = draw_spot(file, visual_settings=visual_settings, show=False)
        image_src = f"{self._report_tmp}image_{self._images_count}.png"
        fig.savefig(image_src)
        plt.close(fig)
        self.add_image(image_src, f'SPOT диаграмма для полей с углами падения: {angles}.',
                       size=(visual_settings.width, visual_settings.height))

    def _vignetting_tables(self, file: ResultFile):
        self.add_paragraph(f"\tИнформация о параметрах виньетирования относительно координат зрачка представлена"
                           f" в следующих трёх таблицах {self._tables_count + 1} - {self._images_count + 1 + 3}.",
                           alignment=WD_PARAGRAPH_ALIGNMENT.JUSTIFY)

        data = []
        vig1 = file.vignetting_info(1)[-1]
        vig2 = file.vignetting_info(5)[-1]
        vig3 = file.vignetting_info(9)[-1]
        wl1, x, y, vig_1 = vig1['wave_id'], vig1['pupil_x'], vig1['pupil_y'], vig1['vignetting']
        wl2, x, y, vig_2 = vig2['wave_id'], vig2['pupil_x'], vig2['pupil_y'], vig2['vignetting']
        wl3, x, y, vig_3 = vig3['wave_id'], vig3['pupil_x'], vig3['pupil_y'], vig3['vignetting']

        for WAVE_ID, vig in zip((wl1, wl2, wl3), (vig_1, vig_2, vig_3)):
            data.clear()
            wl = file.wavelengths[WAVE_ID]
            for index, value in enumerate(vig.flat):
                row, col = divmod(index, x.size)
                if col == 0:
                    data.append(f'{y[row]:.2f}')
                    data.append(f'{value:.2f}')
                else:
                    data.append(f'{value:.2f}')
            header = ['PY/PX']
            header.extend([f'{xi:.2f}' for xi in x])
            # self._docx.add_section(start_type=WD_SECTION_START.NEW_PAGE)
            # self._docx.sections[-1].orientation = WD_ORIENTATION.LANDSCAPE
            # self._docx.sections[-1].page_width = self._docx.sections[0].page_height
            # self._docx.sections[-1].page_height = self._docx.sections[0].page_width
            self.add_table(headers=header,
                           description=f'Параметры веньетирования в зависимости от нормализованных координат входного'
                                       f' зрачка диаграммы при длине волны {wl}мКм',
                           data=data, font_size=12)
        section = self._docx.sections[-1]
        section.orientation = WD_ORIENTATION.LANDSCAPE

    def _polarization_tables(self, file: ResultFile):
        self.add_paragraph(f"\tИнформация о параметрах интенсивности  относительно координат зрачка представлена"
                           f" в следующих трёх таблицах {self._tables_count + 1} - {self._images_count + 1 + 3}.",
                           alignment=WD_PARAGRAPH_ALIGNMENT.JUSTIFY)
        data = []
        x, y, vig_1 = file.vignetting_info(1)
        x, y, vig_2 = file.vignetting_info(5)
        x, y, vig_3 = file.vignetting_info(9)
        for WAVE_ID, vig in zip((1, 5, 9), (vig_1, vig_2, vig_3)):
            data.clear()
            wl = file.wavelengths[WAVE_ID]
            for index, value in enumerate(vig.flat):
                row, col = divmod(index, x.size)
                if col == 0:
                    data.append(f'{y[row]:.2f}')
                    data.append(f'{value:.2f}')
                else:
                    data.append(f'{value:.2f}')
            header = ['PY/PX']
            header.extend([f'{xi:.2f}' for xi in x])
            self.add_table(headers=header,
                           description=f'Параметры веньетирования в зависимости от нормализованных координат входного'
                                       f' зрачка диаграммы при длине волны {wl}мКм',
                           data=data, font_size=12)

    def _spot_diagrams_tables(self, file: ResultFile):
        self.add_paragraph(f"\tДетальная информация о параметрах спот диаграммы представлена в таблицах "
                           f"{self._tables_count + 1} - {self._images_count + 1 + len(file.wavelengths)}."
                           f"Параметры включают в себя координаты геометрического центра лучей на плоскости изображения"
                           f" в зависимости от угла падения и длины волны, "
                           f"средний геометрический и средне квадратичный радиусы.",
                           alignment=WD_PARAGRAPH_ALIGNMENT.JUSTIFY)
        data = []
        for WAVE_ID in range(len(file.wavelengths)):
            data.clear()
            wl = file.wavelengths[WAVE_ID]
            spots = file.get_spot(field_id=-1, wave_id=WAVE_ID + 1)
            for v in spots:
                data.append(f'{v.FIELD_ID}')
                data.append(f'{file.fields[v.FIELD_ID].FLDX:.3f}')
                data.append(f'{file.fields[v.FIELD_ID].FLDY:.3f}')
                data.append(f'{v.CENTER.x:.3f}')
                data.append(f'{v.CENTER.y:.3f}')
                data.append(f'{v.R_GEO:.3f}')
                data.append(f'{v.R_RMS:.3f}')
            self.add_table(headers=('Номер поля', 'Угол по Х', 'Угол по Y', 'X-центр', 'Y-центр', 'R-AVG', 'R-RMS'),
                           description=f'Параметры спот диаграммы для каждого из полей при длине волны {wl}мКм. '
                                       f'Углы имеют размерность градусов, координаты и радиусы - миллиметры',
                           data=data, font_size=12)

    def _spot_center_positions_interp(self, file: ResultFile):
        """
        #######################################
        ######     ПРЯМАЯ ЗАВИСИМОСТЬ    ######
        #######################################
        """
        # self.add_paragraph(
        #     f"\n\tПрямая зависимость положения геометрического центра спот диаграммы на"
        #     f" плоскости изображения в виде полиномиального разложения "
        #     f"представлена в таблицах {self._tables_count + 1} - {self._tables_count + 2}, соответственно.",
        #     alignment=WD_PARAGRAPH_ALIGNMENT.JUSTIFY)
#
        # data = []
        # for coord in file.cords:
        #     angles = np.linspace(coord.AX_MIN, coord.AX_MAX, coord.N_ANGLES_X)
        #     data.append(f'{file.wavelengths[coord.WAVE_ID - 1]:.3f}')
        #     data.extend([f'{v:.3E}' for v in poly_regression(angles, coord.x_slice, 6).flat])
        # self.add_table(headers=('WL,[мкм]', 'X^0', 'X^1', 'X^2', 'X^3', 'X^4', 'X^5'),
        #                description=f'Полиномиальное разложение прямой зависимости положения геометрического центра '
        #                            f'спот диаграммы на '
        #                            f'плоскости относительно оси х. Коэффициенты имеют размерности соответственно'
        #                            f' mm, mm^0, mm^-1, mm^-2, mm^-3.', data=data, font_size=12)
        # data.clear()
        # for coord in file.cords:
        #     angles = np.linspace(coord.AX_MIN, coord.AX_MAX, coord.N_ANGLES_X)
        #     data.append(f'{file.wavelengths[coord.WAVE_ID - 1]:.3f}')
        #     data.extend(list(map(lambda v: f'{v:.3E}', poly_regression(angles, coord.y_slice, 6).flat)))
        # self.add_table(headers=('WL,[мкм]', 'Y^0', 'Y^1', 'Y^2', 'Y^3', 'Y^4', 'Y^5'),
        #                description=f'Полиномиальное разложение прямой зависимости положения геометрического центра '
        #                            f'спот диаграммы на '
        #                            f'плоскости относительно оси y. Коэффициенты имеют размерности соответственно'
        #                            f' mm, mm^0, mm^-1, mm^-2, mm^-3.', data=data, font_size=12)
        # self.add_paragraph(
        #     f"\n\tПрямая зависимость положения геометрического центра спот диаграммы на плоскости изображения "
        #     f"в виде линейного разложения "
        #     f"представлено в таблицах {self._tables_count + 1} - {self._tables_count + 2}. Коэффициенты имеют "
        #     f"размерности соответственно mm, mm^0.", alignment=WD_PARAGRAPH_ALIGNMENT.JUSTIFY)
        # data.clear()
        # for coord in file.cords:
        #     angles_x = np.linspace(coord.AX_MIN, coord.AX_MAX, coord.N_ANGLES_X)
        #     angles_y = np.linspace(coord.AX_MIN, coord.AY_MAX, coord.N_ANGLES_Y)
        #     data.append(f'{file.wavelengths[coord.WAVE_ID - 1]:.3f}')
        #     data.extend(list(map(lambda v: f'{v:.3E}', linear_regression(angles_x, coord.x_slice))))
        #     data.extend(list(map(lambda v: f'{v:.3E}', linear_regression(angles_y, coord.y_slice))))
        # self.add_table(headers=('WL,[мкм]', 'kx', 'bx', 'ky', 'by'),
        #                description=f'Линейное разложение прямых зависимостей положения центра спот диаграммы на  '
        #                            f'плоскости относительно осей x и y.', data=data, font_size=12)
        """
        #######################################
        ######    ОБРАТНАЯ ЗАВИСИМОСТЬ   ######
        #######################################
        """
        self.add_paragraph(
            f"\n\tЗависимость угла падения поля от положения геометрического центра спот диаграммы на"
            f" плоскости изображения в виде полиномиального разложения "
            f"представлена в таблицах {self._tables_count + 1} - {self._tables_count + 2}, соответственно.",
            alignment=WD_PARAGRAPH_ALIGNMENT.JUSTIFY)

        data = []
        for coord in file.cords:
            angles = np.linspace(coord.AX_MIN, coord.AX_MAX, coord.N_ANGLES_X)
            data.append(f'{file.wavelengths[coord.WAVE_ID - 1]:.3f}')
            data.extend([f'{v:.3f}' for v in poly_regression(coord.x_slice, angles, 6).flat])
        self.add_table(headers=('WL,[мкм]', 'X^0', 'X^1', 'X^2', 'X^3', 'X^4', 'X^5'),
                       description=f'Полиномиальное разложение зависимости угла падения поля от'
                                   f' положения геометрического центра спот диаграммы на '
                                   f'плоскости относительно оси х. Коэффициенты имеют размерности соответственно'
                                   f' mm, mm^0, mm^-1, mm^-2, mm^-3', data=data)  # , font_size=12)
        data.clear()
        for coord in file.cords:
            angles = np.linspace(coord.AX_MIN, coord.AX_MAX, coord.N_ANGLES_X)
            data.append(f'{file.wavelengths[coord.WAVE_ID - 1]:.3f}')
            data.extend(list(map(lambda v: f'{v:.3f}', poly_regression(coord.y_slice, angles, 6).flat)))
        self.add_table(headers=('WL,[мкм]', 'Y^0', 'Y^1', 'Y^2', 'Y^3', 'Y^4', 'Y^5'),
                       description=f'Полиномиальное разложение зависимости угла падения поля от'
                                   f' положения геометрического центра спот диаграммы на '
                                   f'плоскости относительно оси y. Коэффициенты имеют размерности соответственно'
                                   f' mm, mm^0, mm^-1, mm^-2, mm^-3', data=data)  # , font_size=12)

        # self.add_paragraph(
        #     f"\tОбратная зависимость положения геометрического центра спот диаграммы на плоскости изображения "
        #     f"в виде линейного разложения "
        #     f"представлено в таблицах {self._tables_count + 1} - {self._tables_count + 2}. Коэффициенты имеют "
        #     f"размерности соответственно mm, mm^0.", alignment=WD_PARAGRAPH_ALIGNMENT.JUSTIFY)
        # data.clear()
        # for coord in file.cords:
        #     angles_x = np.linspace(coord.AX_MIN, coord.AX_MAX, coord.N_ANGLES_X)
        #     angles_y = np.linspace(coord.AX_MIN, coord.AY_MAX, coord.N_ANGLES_Y)
        #     data.append(f'{file.wavelengths[coord.WAVE_ID - 1]:.3f}')
        #     data.extend(list(map(lambda v: f'{v:.3E}', linear_regression(coord.x_slice, angles_x))))
        #     data.extend(list(map(lambda v: f'{v:.3E}', linear_regression(coord.y_slice, angles_y))))
        # self.add_table(headers=('WL,[мкм]', 'kx', 'bx', 'ky', 'by'),
        #                description=f'Линейное разложение обратных зависимостей положения центра спот диаграммы на  '
        #                            f'плоскости относительно осей x и y.', data=data, font_size=12)

    def _spot_center_positions_interp_images(self, file: ResultFile):
        self.add_paragraph(
            f"\n\tНа рисунках {self._images_count + 1} и {self._images_count + 2} представлены прямые зависимости"
            f" положения геометрического центра спот диаграммы на плоскости изображения от угла падения поля.\n",
            alignment=WD_PARAGRAPH_ALIGNMENT.JUSTIFY)

        visual_settings = ResultVisualSettings()
        visual_settings.subplots  = (1, 1)
        visual_settings.width = 12
        visual_settings.height = 12
        visual_settings.font_size = 12
        visual_settings.bounds = (0.2, 0.2, 0.8, 0.8)
        visual_settings.h_space = 0.25
        visual_settings.w_space = 0.25

        fig = draw_pos_from_angle(file, visual_settings=visual_settings, direction='x', show=False)
        image_src = f"{self._report_tmp}image_{self._images_count}.png"
        fig.savefig(image_src)
        plt.close(fig)
        self.add_image(image_src, 'Зависимость положения центра спот диаграммы от угла падения X в плоскости ZoY',
                       size=(visual_settings.width, visual_settings.height))

        fig = draw_pos_from_angle(file, visual_settings=visual_settings, direction='y', show=False)
        image_src = f"{self._report_tmp}image_{self._images_count}.png"
        fig.savefig(image_src)
        plt.close(fig)
        self.add_image(image_src, 'Зависимость положения центра спот диаграммы от угла падения Y в плоскости ZoX',
                       size=(visual_settings.width, visual_settings.height))

    def _mtf_diagrams(self, file: ResultFile):
        visual_settings = ResultVisualSettings()
        visual_settings.subplots = (1, 1)
        visual_settings.width = 17
        visual_settings.height = 17
        visual_settings.font_size = 12
        fig = draw_mtf(file, visual_settings=visual_settings, show=False)
        image_src = f"{self._report_tmp}image_{self._images_count}.png"
        fig.savefig(image_src)
        plt.close(fig)
        self.add_paragraph(
            f"\n\tЧастотно контрастная характеристика схемы представлена на рисунке {self._images_count + 1}.\n",
            alignment=WD_PARAGRAPH_ALIGNMENT.JUSTIFY)
        self.add_image(image_src, 'Частотно контрастная характеристика для всех полей, усреднённая по длинам волн',
                       size=(visual_settings.width, visual_settings.height))

    def _psf_diagrams(self, file: ResultFile):
        visual_settings = ResultVisualSettings()
        visual_settings.subplots = (2, 2)
        visual_settings.width = 17
        visual_settings.height = 17
        visual_settings.font_size = 12
        visual_settings.bounds = (0.05, 0.05, 0.95, 0.95)
        visual_settings.h_space = 0.125
        visual_settings.w_space = 0.125
        angles = ",".join(f" {{ x = {file.fields[fid + 1].FLDX:.3}, y = {file.fields[fid + 1].FLDY:.3} }}"
                          for fid in range(visual_settings.subplots_count))
        self.add_paragraph(
            f"\n\tНа рисунке {self._images_count + 1} "
            f"представлены функции рассеяния точки для полей схемы со следующими углами падения: {angles}.\n",
            alignment=WD_PARAGRAPH_ALIGNMENT.JUSTIFY)

        fig = draw_psf(file, visual_settings=visual_settings, show=False)
        image_src = f"{self._report_tmp}image_{self._images_count}.png"
        fig.savefig(image_src)
        plt.close(fig)
        self.add_image(image_src, f'Функции рассеяния точки исследуемой схемы для полей с углами падения: {angles}',
                       size=(visual_settings.width, visual_settings.height))

    def _psf_diagrams_tables(self, file: ResultFile):
        self.add_paragraph(f"\n\tПодробная информация о параметрах функций рассеяния точек схемы представлена в таблицах "
                           f"{self._tables_count + 1} - {self._images_count + 1 + len(file.wavelengths)},"
                           f" В таблицах представлены координаты максимума интенсивности функции рассеяния точки и "
                           f"относительная интенсивность на плоскости изображения от угла падения поля и длины волны. "
                           f"Углы выражены в градусах, координаты в миллиметрах, "
                           f"интенсивность безразмерна и нормирована на единицу.\n",
                           alignment=WD_PARAGRAPH_ALIGNMENT.JUSTIFY)
        data = []
        # psf_x = file.x_image_pos_per_angle_from_psf
        # psf_y = file.y_image_pos_per_angle_from_psf
        fields_psf = [[(v.center_world_space, v.relative_intensity) for v in file.psf_s if v.WAVE_ID == wave_id + 1]
                      for wave_id in range(file.n_wave_lengths)]

        for wl_index, field_xy in enumerate(fields_psf):
            data.clear()
            wl = file.wavelengths[wl_index]
            for field_id, ((x, y), i) in enumerate(field_xy):
                data.append(f'{field_id + 1}')
                data.append(f'{file.fields[field_id + 1].FLDX:.3f}')
                data.append(f'{file.fields[field_id + 1].FLDY:.3f}')
                data.append(f'{x:.3f}')
                data.append(f'{y:.3f}')
                data.append(f'{i:.3f}')
            self.add_table(headers=('Номер поля', 'Угол по Х', 'Угол по Y', 'X-центр', 'Y-центр', 'Интенсивность'),
                           description=f'Параметры функций рассеяния точки схемы для каждого из полей при длине волны '
                                       f'{wl} мКм', data=data)  # ,, font_size=12)

    def _psf_center_positions_interp(self, file: ResultFile):
        """
        #######################################
        ######     ПРЯМАЯ ЗАВИСИМОСТЬ    ######
        #######################################
        """
        # self.add_paragraph(
        #     f"\n\tПрямая зависимость положения максимума интенсивности функции рассеяния точки на"
        #     f" плоскости изображения в виде полиномиального разложения "
        #     f"представлена в таблицах {self._tables_count + 1} - {self._tables_count + 2}, соответственно.\n",
        #     alignment=WD_PARAGRAPH_ALIGNMENT.JUSTIFY)

        # data = []
        # angles_x = file.fields_angles_x
        # angles_y = file.fields_angles_y
        # x_cords = file.x_image_pos_per_angle_from_psf
        # y_cords = file.x_image_pos_per_angle_from_psf
        # for wave_id, psf_cords in enumerate(x_cords):
        #     data.append(f'{file.wavelengths[wave_id]:.3f}')
        #     data.extend([f'{v:.3E}' for v in poly_regression(angles_x, psf_cords, 6).flat])
        # self.add_table(headers=('WL,[мкм]', 'X^0', 'X^1', 'X^2', 'X^3', 'X^4', 'X^5'),
        #                description=f'Полиномиальное разложение прямой зависимости положения максимума интенсивности'
        #                            f' функции рассеяния точки на '
        #                            f'плоскости относительно оси х. Коэффициенты имеют размерности соответственно'
        #                            f' mm, mm^0, mm^-1, mm^-2, mm^-3.', data=data, font_size=12)
        # data.clear()
        # for wave_id, psf_cords in enumerate(y_cords):
        #     data.append(f'{file.wavelengths[wave_id]:.3f}')
        #     data.extend(list(map(lambda v: f'{v:.3E}', poly_regression(angles_y, psf_cords, 6).flat)))
        # self.add_table(headers=('WL,[мкм]', 'Y^0', 'Y^1', 'Y^2', 'Y^3', 'Y^4', 'Y^5'),
        #                description=f'Полиномиальное разложение прямой зависимости положения максимума интенсивности'
        #                            f' функции рассеяния точки на '
        #                            f'плоскости относительно оси y. Коэффициенты имеют размерности соответственно'
        #                            f' mm, mm^0, mm^-1, mm^-2, mm^-3.', data=data, font_size=12)
        # self.add_paragraph(
        #     f"\n\tПрямая зависимость положения максимума интенсивности функции рассеяния точки на плоскости изображения "
        #     f"в виде линейного разложения "
        #     f"представлено в таблицах {self._tables_count + 1} - {self._tables_count + 2}. Коэффициенты имеют "
        #     f"размерности соответственно mm, mm^0.\n", alignment=WD_PARAGRAPH_ALIGNMENT.JUSTIFY)
        # data.clear()
        # for wave_id, (x, y) in enumerate(zip(x_cords, y_cords)):
        #     data.append(f'{file.wavelengths[wave_id]:.3f}')
        #     data.extend(list(map(lambda v: f'{v:.3E}', linear_regression(angles_x, x))))
        #     data.extend(list(map(lambda v: f'{v:.3E}', linear_regression(angles_y, y))))
        # self.add_table(headers=('WL,[мкм]', 'kx', 'bx', 'ky', 'by'),
        #                description=f'Линейное разложение прямых зависимостей положения максимума интенсивности'
        #                            f' функции рассеяния точки на плоскости относительно осей x и y.',
        #                data=data, font_size=12)
        """
        #######################################
        ######    ОБРАТНАЯ ЗАВИСИМОСТЬ   ######
        #######################################
        """
        self.add_paragraph(
            f"\n\tЗависимость угла падения поля от положения максимума интенсивности функции рассеяния точки на"
            f" плоскости изображения в виде полиномиального разложения "
            f"представлена в таблицах {self._tables_count + 1} - {self._tables_count + 2}, соответственно.\n",
            alignment=WD_PARAGRAPH_ALIGNMENT.JUSTIFY)

        data = []
        angles_x = file.fields_angles_x
        angles_y = file.fields_angles_y
        x_cords = file.x_image_pos_per_angle_from_psf
        y_cords = file.x_image_pos_per_angle_from_psf
        for wave_id, psf_cords in enumerate(x_cords):
            data.append(f'{file.wavelengths[wave_id]:.3f}')
            if psf_cords.size != 0:
                length = min(psf_cords.size, angles_x.size)
                data.extend([f'{v:.3f}' for v in poly_regression(psf_cords[0:length], angles_x[0:length], 6).flat])
            else:
                data.extend([f'NAN' for _ in range(6)])

        self.add_table(headers=('WL,[мкм]', 'X^0', 'X^1', 'X^2', 'X^3', 'X^4', 'X^5'),
                       description=f'Полиномиальное разложение зависимости угла падения поля от положения'
                                   f' максимума интенсивности функции рассеяния точки на '
                                   f'плоскости относительно оси х. Коэффициенты имеют размерности соответственно'
                                   f' mm, mm^0, mm^-1, mm^-2, mm^-3', data=data)  # ,, font_size=12)
        data.clear()
        for wave_id, psf_cords in enumerate(y_cords):
            data.append(f'{file.wavelengths[wave_id]:.3f}')
            if psf_cords.size != 0:
                length = min(psf_cords.size, angles_x.size)
                data.extend([f'{v:.3f}' for v in poly_regression(psf_cords[0:length], angles_x[0:length], 6).flat])
            else:
                data.extend([f'NAN' for _ in range(6)])
        self.add_table(headers=('WL,[мкм]', 'Y^0', 'Y^1', 'Y^2', 'Y^3', 'Y^4', 'Y^5'),
                       description=f'Полиномиальное разложение зависимости угла падения поля от положения'
                                   f' максимума интенсивности функции рассеяния точки на '
                                   f'плоскости относительно оси y. Коэффициенты имеют размерности соответственно'
                                   f' mm, mm^0, mm^-1, mm^-2, mm^-3', data=data)  # ,, font_size=12)
        # self.add_paragraph(
        #     f"\n\tОбратная зависимость положения максимума интенсивности функции рассеяния точки на плоскости изображения "
        #     f"в виде линейного разложения "
        #     f"представлено в таблицах {self._tables_count + 1} - {self._tables_count + 2}. Коэффициенты имеют "
        #     f"размерности соответственно mm, mm^0.\n", alignment=WD_PARAGRAPH_ALIGNMENT.JUSTIFY)
        # data.clear()
        # for wave_id, (x, y) in enumerate(zip(x_cords, y_cords)):
        #     data.append(f'{file.wavelengths[wave_id]:.3f}')
        #     data.extend(list(map(lambda v: f'{v:.3E}', linear_regression(x, angles_x))))
        #     data.extend(list(map(lambda v: f'{v:.3E}', linear_regression(y, angles_y))))
        # self.add_table(headers=('WL,[мкм]', 'kx', 'bx', 'ky', 'by'),
        #                description=f'Линейное разложение обратных зависимостей положения максимума интенсивности'
        #                            f' функции рассеяния точки на плоскости относительно осей x и y.',
        #                data=data, font_size=12)

    def _psf_center_positions_interp_images(self, file: ResultFile):
        self.add_paragraph(
            f"\n\tНа рисунках {self._images_count + 1} и {self._images_count + 2} представлены прямые зависимости положения"
            f" центра спот диаграммы на плоскости изображения от угла падения поля.\n",
            alignment=WD_PARAGRAPH_ALIGNMENT.JUSTIFY)

        visual_settings = ResultVisualSettings()
        visual_settings.subplots = (1, 1)
        visual_settings.width = 12
        visual_settings.height = 12
        visual_settings.font_size = 12
        visual_settings.bounds = (0.2, 0.2, 0.8, 0.8)
        visual_settings.h_space = 0.25
        visual_settings.w_space = 0.25
        fig = draw_pos_from_angle(file, visual_settings=visual_settings, direction='x', show=False)
        image_src = f"{self._report_tmp}image_{self._images_count}.png"
        fig.savefig(image_src)
        plt.close(fig)
        self.add_image(image_src, 'Зависимость х-координаты положения максимума интенсивности функции'
                                  ' рассеяния точки от угла падения X в плоскости ZoY',
                       size=(visual_settings.width, visual_settings.height))

        fig = draw_pos_from_angle(file, visual_settings=visual_settings, direction='y', show=False)
        image_src = f"{self._report_tmp}image_{self._images_count}.png"
        fig.savefig(image_src)
        plt.close(fig)
        self.add_image(image_src, 'Зависимость y-координаты положения максимума интенсивности функции'
                                  ' рассеяния точки от угла падения Y в плоскости ZoX',
                       size=(visual_settings.width, visual_settings.height))

    def update(self, file: ResultFile, show_main_info: bool = False):
        self._report_tmp = os.path.curdir + '\\report_tmp\\'
        if not os.path.isdir(self._report_tmp):
            os.mkdir(self._report_tmp)
        self.add_paragraph(f"Результаты моделирования для {file.scheme_src}\n", bold=True, font_size=14)
        if show_main_info:
            self._basic_params(file)
        self._mtf_diagrams(file)
        # self._vignetting_tables(file)
        self._spot_diagrams(file)
        # self._spot_center_positions_interp_images(file)
        self._spot_diagrams_tables(file)
        self._spot_center_positions_interp(file)
        #########################################
        self._psf_diagrams(file)
        self._psf_diagrams_tables(file)
        # self._psf_center_positions_interp_images(file)
        self._psf_diagrams_tables(file)
        self._psf_center_positions_interp(file)
        self.keep_table_on_one_page()

