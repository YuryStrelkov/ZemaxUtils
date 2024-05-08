from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import Cm
from docx.shared import Pt
from docx import Document

image_count = 0
tables_count = 0


def add_image(document, src: str, description: str = ""):
    global image_count
    image_count += 1

    img = document.add_picture(src, width=Cm(12.0))
    last_paragraph = document.paragraphs[-1]
    last_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    p = document.add_paragraph()
    p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    if description == "":
        f_name = '.'.join(v for v in src.split("\\")[-1].split(".")[:-1])
        if f_name.startswith("psf"):
            description = f"Результат моделирования PSF по всем полям для длины волны {f_name.split(' ')[-1]} мкм."
        elif f_name.startswith("psf_cross"):
            description = f"Результат моделирования сечения PSF плоскостями XoZ и YoZ" \
                          f" по всем длинам волн для поля с номером {f_name.split(' ')[-1]}."
        elif f_name.startswith("mtf"):
            description = f"Результат моделирования MTF по всем длинам волн для всех полей."
        elif f_name.startswith("spot"):
            description = f"Результат моделирования SPOT диаграммы по всем длинам волн для всех полей."
        else:
            description = "Нет описания к рисунку."

    t = p.add_run(f"Pис. {image_count}. {description}")
    t.font.name = 'Times New Roman'
    t.font.size = Pt(12)
    t.italic = True
    # document.add_page_break()


def build_table(document, headers, data, description: str = "", format_provider=lambda s: str(s)):
    global tables_count
    tables_count += 1
    table = document.add_table(rows=len(data) // len(headers) + 1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    header_cells = table.rows[0].cells
    for cell, header in zip(header_cells, headers):
        cell.text = str(header)
        font = cell.paragraphs[0].runs[0].font
        font.bold = True
        font.name = 'Times New Roman'
        font.size = Pt(14)
        cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    for record_index, record in enumerate(data):
        row, col = divmod(record_index, len(headers))
        cell = table.rows[row + 1].cells[col]
        cell.text = format_provider(record)
        font = cell.paragraphs[0].runs[0].font
        font.name = 'Times New Roman'
        font.size = Pt(14)

    add_paragraph(document,
                  text=f"Таблица. {tables_count}. {'Нет описания к таблице.' if description == '' else description}",
                  italic=True)


def add_paragraph(document, text: str = 'Paragraph', font: str = 'Times New Roman', font_size: int = 14,
                  alignment=WD_PARAGRAPH_ALIGNMENT.CENTER, italic: bool = False):
    p = document.add_paragraph()
    p.paragraph_format.alignment = alignment
    t = p.add_run(text)
    t.font.name = font
    t.font.size = Pt(font_size)
    t.italic = italic
    return document


if __name__ == "__main__":
    directory = "D:\\Github\\ZemaxUtils\\"
    document = Document()
    build_table(document, ('A', 'B', 'C'), [1, 2, 3, 4, 5, 6, 7, 8, 9], "Kurwa table")
    document.save(directory + 'table_ex.docx')

    # directories = [directory + f + "\\" for f in os.listdir(directory) if os.path.isdir(directory + f)]
    # sep = "\\"
    # for d in directories:
    #     p = document.add_paragraph()
    #     t = p.add_run(f"Результаты моделирования для \"{d.split(sep)[-2]}\"")
    #     t.font.name = 'Times New Roman'
    #     t.font.size = Pt(14)
    #     t.bold = True
    #     images_1 = [build_image(document, d + f) for f in os.listdir(d) if f.endswith('png')]
    #     images_2 = [build_image(document, d + "PSF\\" + f) for f in os.listdir(d + "PSF\\") if f.endswith('png')]
    #     images_3 = [build_image(document, d + "PSF_CROSS_SECTION\\" + f) for f in os.listdir(d + "PSF_CROSS_SECTION\\") if f.endswith('png')]

